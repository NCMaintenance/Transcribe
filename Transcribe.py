import streamlit as st
import google.generativeai as genai
import json
import os
from datetime import datetime
from docx import Document
import io
import tempfile
import re
from moviepy.editor import VideoFileClip
from pydub import AudioSegment

# --- Utility to prettify keys ---
def prettify_key(key):
    key = key.replace('_', ' ')
    key = re.sub(r'([a-z])([A-Z])', r'\1 \2', key)
    return key.title() + ":"

# --- Configure Gemini API ---
genai.configure(api_key=st.secrets["GEMINI_API_KEY"])
model = genai.GenerativeModel(model_name='gemini-2.0-flash-exp')

st.set_page_config(page_title="Dr. Scribe", layout="wide")

# --- Password protection ---
if "password_verified" not in st.session_state:
    st.session_state.password_verified = False

if not st.session_state.password_verified:
    user_password = st.text_input("Enter password to access Dr. Scribe:", type="password", key="password_input")
    submit_button = st.button("Submit", key="submit_pwd")
    if submit_button:
        if user_password == st.secrets["password"]:
            st.session_state.password_verified = True
            st.rerun()
        else:
            st.warning("Incorrect password. Please try again.")
    st.stop()

# --- Sidebar ---
with st.sidebar:
    st.image("https://www.ehealthireland.ie/media/k1app1wt/hse-logo-black-png.png", width=200)
    st.title("🩺 Dr. Scribe")
    if st.button("Created by Dave Maher"):
        st.sidebar.write("This application intellectual property belongs to Dave Maher.")

# --- Main UI ---
st.title("🩺 Dr. Scribe")
st.markdown("### 📤 Upload or Record Doctor–Patient Audio")

mode = st.radio("Choose input method:", ["Upload audio/video file", "Record using microphone"])

audio_bytes = None
audio_format = "audio/mp3"

if mode == "Upload audio/video file":
    uploaded_audio = st.file_uploader("Upload an audio or video file (WAV, MP3, M4A, MP4)", type=["wav", "mp3", "m4a", "mp4"])
    if uploaded_audio:
        st.audio(uploaded_audio)
        audio_bytes = uploaded_audio

elif mode == "Record using microphone":
    recorded_audio = st.audio_input("🎙️ Click the microphone to record")
    if recorded_audio:
        st.audio(recorded_audio)
        audio_bytes = recorded_audio

# --- Transcription and Analysis ---
if audio_bytes and st.button("🧠 Transcribe & Analyse"):
    with st.spinner("Processing with Gemini..."):
        if hasattr(audio_bytes, "read"):
            audio_bytes = audio_bytes.read()

        # Save uploaded file
        original_suffix = uploaded_audio.name.split(".")[-1].lower() if mode == "Upload audio/video file" else ".wav"
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{original_suffix}") as temp_input:
            temp_input.write(audio_bytes)
            temp_input_path = temp_input.name

        try:
            # Convert to MP3
            with tempfile.NamedTemporaryFile(delete=False, suffix=".mp3") as temp_mp3:
                temp_mp3_path = temp_mp3.name

            if original_suffix == "mp4":
                video = VideoFileClip(temp_input_path)
                audio_path_wav = temp_input_path.replace(".mp4", ".wav")
                video.audio.write_audiofile(audio_path_wav)
                sound = AudioSegment.from_wav(audio_path_wav)
                sound.export(temp_mp3_path, format="mp3")
                video.close()
                os.remove(audio_path_wav)
            else:
                sound = AudioSegment.from_file(temp_input_path)
                sound.export(temp_mp3_path, format="mp3")

            # Upload and process
            audio_file = genai.upload_file(path=temp_mp3_path)
            prompt = (
                "You are a medical transcriptionist. Transcribe the following doctor–patient consultation. "
                "Label speakers as 'Doctor:' or 'Patient:' where possible."
            )
            result = model.generate_content([prompt, audio_file], request_options={"timeout": 600})
            transcript = result.text
            genai.delete_file(audio_file.name)
            st.session_state["transcript"] = transcript
            st.success("Transcript generated successfully.")

        finally:
            os.remove(temp_input_path)
            if os.path.exists(temp_mp3_path):
                os.remove(temp_mp3_path)

# --- Display Transcript ---
if "transcript" in st.session_state:
    st.markdown("## 📄 Transcript")
    st.text_area("Transcript", st.session_state["transcript"], height=300)

    if st.button("📊 Summarise Transcript"):
        with st.spinner("Generating structured and narrative summaries..."):
            # Structured Summary Prompt
            prompt_structured = f"""
You are a medical scribe. Extract key details from this doctor–patient transcript and return JSON with:
- patientName
- dateOfVisit
- chiefComplaint
- historyPresentIllness
- pastMedicalHistory
- medications
- allergies
- reviewOfSystems
- physicalExam
- assessment
- plan
- followUp
If not mentioned, use "Not mentioned".
Transcript:
{st.session_state['transcript']}
            """
            response1 = model.generate_content(prompt_structured)

            json_match = re.search(r"\{.*\}", response1.text, re.DOTALL)
            if json_match:
                try:
                    structured = json.loads(json_match.group())
                except json.JSONDecodeError as e:
                    st.error("❌ JSON parse error.")
                    st.code(json_match.group(), language="json")
                    raise e
            else:
                st.error("❌ No valid JSON found.")
                st.code(response1.text)
                raise ValueError("Invalid JSON.")

            # Narrative Summary Prompt
            prompt_narrative = f"""
Summarise the transcript into a coherent, professional doctor’s narrative summary using appropriate medical language.
Transcript:
{st.session_state['transcript']}
            """
            response2 = model.generate_content(prompt_narrative)
            st.session_state["structured"] = structured
            st.session_state["narrative"] = response2.text
            st.success("Summaries generated.")

# --- DOCX Export ---
def create_docx(content, kind="structured"):
    doc = Document()
    if kind == "structured":
        doc.add_heading("Structured Medical Summary", level=1)
        for key, val in content.items():
            doc.add_heading(prettify_key(key), level=2)
            doc.add_paragraph(val)
    else:
        doc.add_heading("Doctor’s Narrative Summary", level=1)
        doc.add_paragraph(content)
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

# --- Display Results ---
if "structured" in st.session_state and "narrative" in st.session_state:
    st.markdown("## 📑 Structured Summary")
    for k, v in st.session_state["structured"].items():
        st.markdown(f"**{prettify_key(k)}** {v}")

    st.download_button("📥 Download Structured Summary (DOCX)",
        data=create_docx(st.session_state["structured"], "structured"),
        file_name="structured_summary.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    st.markdown("---")
    st.markdown("## 🧑‍⚕️ Doctor’s Narrative Summary")
    st.write(st.session_state["narrative"])

    st.download_button("📥 Download Narrative Summary (DOCX)",
        data=create_docx(st.session_state["narrative"], "narrative"),
        file_name="narrative_summary.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# import streamlit as st
# import google.generativeai as genai
# import json
# import os
# from datetime import datetime
# from docx import Document
# import io
# import tempfile
# import re

# # --- Utility to prettify keys ---
# def prettify_key(key):
#     # Convert snake_case to spaces
#     key = key.replace('_', ' ')
#     # Convert camelCase to spaces (insert space before uppercase that follows lowercase)
#     key = re.sub(r'([a-z])([A-Z])', r'\1 \2', key)
#     # Capitalize each word, add colon
#     return key.title() + ":"

# # --- Configure Gemini API ---
# genai.configure(api_key=st.secrets["GEMINI_API_KEY"])
# model = genai.GenerativeModel(model_name='gemini-2.0-flash-exp')

# st.set_page_config(page_title="Dr. Scribe", layout="wide")

# # --- Password protection ---
# if "password_verified" not in st.session_state:
#     st.session_state.password_verified = False

# if not st.session_state.password_verified:
#     user_password = st.text_input("Enter password to access Dr. Scribe:", type="password", key="password_input")
#     submit_button = st.button("Submit", key="submit_pwd")
#     if submit_button:
#         password = st.secrets["password"]
#         if user_password == password:
#             st.session_state.password_verified = True
#             st.rerun()
#         else:
#             st.warning("Incorrect password. Please try again.")
#     st.stop()

# # --- Sidebar with logo, title, button ---
# with st.sidebar:
#     st.image("https://www.ehealthireland.ie/media/k1app1wt/hse-logo-black-png.png", width=200)
#     st.title("🩺 Dr. Scribe")
#     if st.button("Created by Dave Maher"):
#         st.sidebar.write("This application intellectual property belongs to Dave Maher.")

# # --- Main UI ---
# st.title("🩺 Dr. Scribe")
# st.markdown("### 📤 Upload or Record Doctor–Patient Audio")

# # --- Input Method Selection ---
# mode = st.radio("Choose input method:", ["Upload audio file", "Record using microphone"])

# audio_bytes = None
# audio_format = "audio/wav"

# if mode == "Upload audio file":
#     uploaded_audio = st.file_uploader("Upload an audio file (WAV, MP3, M4A)", type=["wav", "mp3", "m4a"])
#     if uploaded_audio:
#         st.audio(uploaded_audio, format=audio_format)
#         audio_bytes = uploaded_audio  # File-like object

# elif mode == "Record using microphone":
#     recorded_audio = st.audio_input("🎙️ Click the microphone to record, then click again to stop and process.")
#     if recorded_audio:
#         st.audio(recorded_audio, format=audio_format)
#         audio_bytes = recorded_audio  # Already in bytes

# # --- Transcription and Analysis ---
# if audio_bytes and st.button("🧠 Transcribe & Analyse"):
#     with st.spinner("Processing with Gemini..."):

#         # Convert to bytes if it's a file-like object
#         if hasattr(audio_bytes, "read"):
#             audio_bytes = audio_bytes.read()

#         # Save to temporary file
#         with tempfile.NamedTemporaryFile(delete=False, suffix=".wav") as tmp_file:
#             tmp_file.write(audio_bytes)
#             tmp_file_path = tmp_file.name

#         try:
#             audio_file = genai.upload_file(path=tmp_file_path)

#             prompt = (
#                 "You are a medical transcriptionist. Transcribe the following doctor–patient consultation. "
#                 "Label speakers as 'Doctor:' or 'Patient:' where possible."
#             )
#             result = model.generate_content([prompt, audio_file], request_options={"timeout": 600})
#             transcript = result.text
#             genai.delete_file(audio_file.name)

#             st.session_state["transcript"] = transcript
#             st.success("Transcript generated successfully.")
#         finally:
#             os.remove(tmp_file_path)

# # --- Display Transcript ---
# if "transcript" in st.session_state:
#     st.markdown("## 📄 Transcript")
#     st.text_area("Transcript", st.session_state["transcript"], height=300)

#     if st.button("📊 Summarise Transcript"):
#         with st.spinner("Generating structured and narrative summaries..."):

#             # --- Structured Summary ---
#             prompt_structured = f"""
# You are a medical scribe. Extract key details from this doctor–patient transcript and return JSON with:
# - patientName
# - dateOfVisit
# - chiefComplaint
# - historyPresentIllness
# - pastMedicalHistory
# - medications
# - allergies
# - reviewOfSystems
# - physicalExam
# - assessment
# - plan
# - followUp

# If not mentioned, use "Not mentioned".
# Transcript:
# {st.session_state['transcript']}
#             """
#             response1 = model.generate_content(prompt_structured)

#             # Parse JSON safely
#             json_match = re.search(r"\{.*\}", response1.text, re.DOTALL)
#             if json_match:
#                 try:
#                     structured = json.loads(json_match.group())
#                 except json.JSONDecodeError as e:
#                     st.error("❌ JSON found but failed to parse. Check formatting.")
#                     st.code(json_match.group(), language="json")
#                     raise e
#             else:
#                 st.error("❌ No valid JSON object found in Gemini's response.")
#                 st.code(response1.text)
#                 raise ValueError("No valid JSON found.")

#             # --- Narrative Summary ---
#             prompt_narrative = f"""
# Summarise the transcript into a coherent, professional doctor’s narrative summary using appropriate medical language.
# Transcript:
# {st.session_state['transcript']}
#             """
#             response2 = model.generate_content(prompt_narrative)
#             narrative = response2.text

#             st.session_state["structured"] = structured
#             st.session_state["narrative"] = narrative
#             st.success("Summaries generated.")

# # --- DOCX Export Function ---
# def create_docx(content, kind="structured"):
#     doc = Document()
#     if kind == "structured":
#         doc.add_heading("Structured Medical Summary", level=1)
#         for key, val in content.items():
#             doc.add_heading(prettify_key(key), level=2)
#             doc.add_paragraph(val)
#     else:
#         doc.add_heading("Doctor’s Narrative Summary", level=1)
#         doc.add_paragraph(content)
#     output = io.BytesIO()
#     doc.save(output)
#     output.seek(0)
#     return output

# # --- Display Summaries and Downloads ---
# if "structured" in st.session_state and "narrative" in st.session_state:
#     st.markdown("## 📑 Structured Summary")
#     for k, v in st.session_state["structured"].items():
#         st.markdown(f"**{prettify_key(k)}** {v}")

#     st.download_button("📥 Download Structured Summary (DOCX)",
#         data=create_docx(st.session_state["structured"], "structured"),
#         file_name="structured_summary.docx",
#         mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

#     st.markdown("---")
#     st.markdown("## 🧑‍⚕️ Doctor’s Narrative Summary")
#     st.write(st.session_state["narrative"])

#     st.download_button("📥 Download Narrative Summary (DOCX)",
#         data=create_docx(st.session_state["narrative"], "narrative"),
#         file_name="narrative_summary.docx",
#         mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
